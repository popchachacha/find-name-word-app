"""
Simplified core processor for Word document analysis.
Focuses on essential functionality without heavy dependencies.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

try:
    from docx import Document
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")


@dataclass
class CharacterStat:
    """Represents the number of times a single character was mentioned."""

    name: str
    count: int


@dataclass
class TableData:
    """Simple representation of a Word table used for re-creation."""

    rows: List[List[str]]


@dataclass
class ProcessingResult:
    """Full snapshot of the processed document."""

    characters: List[str]
    tables: List[TableData]


class DocumentProcessor:
    """Encapsulates processing logic for Word documents."""

    def process(self, source: Path, column_index: int = 1) -> ProcessingResult:
        """
        Read a document and extract characters and table contents.
        
        Args:
            source: Path to the Word document
            column_index: Zero-based index of the column containing character names
            
        Returns:
            ProcessingResult with extracted characters and table data
            
        Raises:
            ValueError: If no characters found in specified column
            FileNotFoundError: If document doesn't exist
        """
        if not source.exists():
            raise FileNotFoundError(f"Document not found: {source}")
            
        if column_index < 0:
            raise ValueError("Column index cannot be negative")

        try:
            document = Document(source)
        except Exception as e:
            raise ValueError(f"Failed to read document: {e}")

        characters: List[str] = []
        tables: List[TableData] = []

        for table_idx, table in enumerate(document.tables):
            table_rows: List[List[str]] = []
            
            for row_idx, row in enumerate(table.rows):
                cell_texts = [cell.text.strip() for cell in row.cells]
                table_rows.append(cell_texts)

                # Extract character name from specified column
                if len(cell_texts) > column_index and cell_texts[column_index]:
                    character_name = cell_texts[column_index]
                    if character_name:  # Only add non-empty names
                        characters.append(character_name)
                        
            tables.append(TableData(rows=table_rows))

        if not characters:
            human_column = column_index + 1
            raise ValueError(
                f"No characters found in column #{human_column} in any table. "
                f"Please check the column index and document structure."
            )

        return ProcessingResult(characters=characters, tables=tables)

    @staticmethod
    def summarise(
        characters: Iterable[str], ignore_case: bool = False
    ) -> List[CharacterStat]:
        """
        Return a sorted list of character stats from the provided names.
        
        Args:
            characters: Iterable of character names
            ignore_case: Whether to merge names regardless of case
            
        Returns:
            List of CharacterStat objects sorted by count (descending)
        """
        if ignore_case:
            display_names: dict[str, str] = {}
            counts: Counter[str] = Counter()
            
            for name in characters:
                if not name:  # Skip empty names
                    continue
                    
                key = name.casefold()
                counts[key] += 1
                # Keep the first encountered name as display name
                display_names.setdefault(key, name)
                
            return [
                CharacterStat(display_names[key], count)
                for key, count in counts.most_common()
            ]

        # Case-sensitive counting
        counts = Counter(char for char in characters if char)  # Skip empty
        return [CharacterStat(name, count) for name, count in counts.most_common()]

    def export_report(
        self,
        result: ProcessingResult,
        output_path: Path,
        minimum_mentions: int = 1,
        ignore_case: bool = False,
    ) -> Path:
        """
        Create a new Word document with the statistics and original tables.
        
        Args:
            result: ProcessingResult to export
            output_path: Path where to save the report
            minimum_mentions: Only include characters with at least this many mentions
            ignore_case: Whether to merge case-insensitive names
            
        Returns:
            Path to the created report file
        """
        # Filter and sort characters
        stats = [
            stat
            for stat in self.summarise(
                result.characters, ignore_case=ignore_case
            )
            if stat.count >= minimum_mentions
        ]

        # Create output document
        output_doc = Document()
        
        # Add title
        title = output_doc.add_heading("Character Frequency Analysis", level=1)
        
        # Add summary statistics
        output_doc.add_heading("Summary", level=2)
        summary_text = f"Total characters found: {len(result.characters)}\n"
        summary_text += f"Unique characters: {len(stats)}\n"
        summary_text += f"Characters with {minimum_mentions}+ mentions: {len(stats)}"
        output_doc.add_paragraph(summary_text)
        
        # Add character list
        output_doc.add_heading("Characters by Frequency", level=2)
        if stats:
            for stat in stats:
                paragraph = output_doc.add_paragraph()
                paragraph.add_run(f"{stat.name}").bold = True
                paragraph.add_run(f" — {stat.count} mentions")
        else:
            output_doc.add_paragraph("No characters meet the minimum frequency criteria.")

        # Add original tables
        if result.tables:
            output_doc.add_heading("Original Tables", level=2)
            for idx, table in enumerate(result.tables, start=1):
                if not table.rows:
                    continue
                    
                output_doc.add_heading(f"Table {idx}", level=3)
                
                # Create table with proper dimensions
                max_cols = max(len(row) for row in table.rows) if table.rows else 0
                if max_cols > 0:
                    new_table = output_doc.add_table(
                        rows=len(table.rows), cols=max_cols
                    )
                    new_table.style = "Table Grid"
                    
                    # Fill table with data
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row):
                            if j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = str(cell)

        # Save document
        output_doc.save(output_path)
        return output_path

    def get_table_preview(self, source: Path, max_rows: int = 5) -> List[List[str]]:
        """
        Get a preview of table structure for UI display.
        
        Args:
            source: Path to Word document
            max_rows: Maximum number of rows to preview per table
            
        Returns:
            List of table previews, each containing up to max_rows
        """
        if not source.exists():
            raise FileNotFoundError(f"Document not found: {source}")
            
        try:
            document = Document(source)
        except Exception as e:
            raise ValueError(f"Failed to read document: {e}")
            
        previews = []
        
        for table in document.tables:
            table_preview = []
            for i, row in enumerate(table.rows[:max_rows]):
                row_data = [cell.text.strip() for cell in row.cells]
                table_preview.append(row_data)
            previews.append(table_preview)
            
        return previews