"""Business logic for analysing Word tables and exporting reports."""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document


@dataclass
class CharacterStat:
    """Represents the number of times a single character was mentioned."""

    name: str
    count: int


@dataclass
class TableData:
    """Simple representation of a Word table used for re-creation."""

    rows: List[List[str]]


@dataclass
class ProcessingResult:
    """Full snapshot of the processed document."""

    characters: List[str]
    tables: List[TableData]


class DocumentProcessor:
    """Encapsulates processing logic for Word documents."""

    def process(self, source: Path, column_index: int = 1) -> ProcessingResult:
        """Read a document and extract characters and table contents."""

        if column_index < 0:
            raise ValueError("Номер столбца не может быть отрицательным")

        document = Document(source)
        characters: List[str] = []
        tables: List[TableData] = []

        for table in document.tables:
            table_rows: List[List[str]] = []
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                table_rows.append(cell_texts)

                if len(cell_texts) > column_index and cell_texts[column_index]:
                    characters.append(cell_texts[column_index])
            tables.append(TableData(rows=table_rows))

        if not characters:
            human_column = column_index + 1
            raise ValueError(
                "Не удалось найти персонажей в указанном столбце"
                f" №{human_column} ни в одной таблице"
            )

        return ProcessingResult(characters=characters, tables=tables)

    @staticmethod
    def summarise(
        characters: Iterable[str], ignore_case: bool = False
    ) -> List[CharacterStat]:
        """Return a sorted list of character stats from the provided names."""

        if ignore_case:
            display_names: dict[str, str] = {}
            counts: Counter[str] = Counter()
            for name in characters:
                key = name.casefold()
                counts[key] += 1
                display_names.setdefault(key, name)
            return [
                CharacterStat(display_names[key], count)
                for key, count in counts.most_common()
            ]

        counts = Counter(characters)
        return [CharacterStat(name, count) for name, count in counts.most_common()]

    def export_report(
        self,
        result: ProcessingResult,
        output_path: Path,
        minimum_mentions: int = 1,
        ignore_case: bool = False,
    ) -> Path:
        """Create a new Word document with the statistics and original tables."""

        stats = [
            stat
            for stat in self.summarise(
                result.characters, ignore_case=ignore_case
            )
            if stat.count >= minimum_mentions
        ]

        output_doc = Document()
        output_doc.add_heading(
            "Количество упоминаний персонажей", level=1
        )
        for stat in stats:
            output_doc.add_paragraph(f"{stat.name} — {stat.count}")

        output_doc.add_paragraph("\nИсходный текст всех таблиц:\n")
        for idx, table in enumerate(result.tables, start=1):
            output_doc.add_paragraph(f"Таблица {idx}:")
            if not table.rows:
                continue
            new_table = output_doc.add_table(
                rows=len(table.rows), cols=max(len(row) for row in table.rows)
            )
            new_table.style = "Table Grid"
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row):
                    new_table.rows[i].cells[j].text = cell

        output_doc.save(output_path)
        return output_path
