from collections import Counter
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document


def process_document(
    input_path,
    output_path,
    target_column_index=1,
    highlight_top=3,
    skip_table_headers=False,
):
    """Обрабатывает один .docx файл и формирует отчёт.

    Args:
        input_path (str | Path): путь к исходному документу.
        output_path (str | Path): путь, по которому нужно сохранить отчёт.
        target_column_index (int): индекс столбца (с нуля), из которого считываются имена.
        highlight_top (int): количество строк лидеров, выделяемых форматированием.
        skip_table_headers (bool): пропускать ли первую строку каждой таблицы как заголовок.
    """

    if target_column_index < 0:
        raise ValueError("Индекс столбца не может быть отрицательным")

    input_doc = Document(input_path)
    characters = []

    for table_idx, table in enumerate(input_doc.tables):
        for row_idx, row in enumerate(table.rows):
            if skip_table_headers and row_idx == 0:
                continue
            try:
                if len(row.cells) > target_column_index:
                    character = row.cells[target_column_index].text.strip()
                    if character:
                        characters.append(character)
                else:
                    print(
                        f"Строка {row_idx + 1} в таблице {table_idx + 1} "
                        "короче выбранного количества столбцов"
                    )
            except IndexError:
                print(
                    f"Ошибка в таблице {table_idx + 1}, строке {row_idx + 1}: "
                    "нет выбранного столбца"
                )

    if not characters:
        raise ValueError(
            "Не удалось найти значения в указанном столбце ни в одной таблице"
        )

    character_counts = Counter(characters)
    sorted_counts = character_counts.most_common()

    output_doc = Document()
    output_doc.add_heading("Количество упоминаний персонажей", level=1)

    summary_table = output_doc.add_table(rows=len(sorted_counts) + 1, cols=2)
    summary_table.style = "Light List"
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Персонаж"
    header_cells[1].text = "Количество"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for idx, (char, count) in enumerate(sorted_counts, start=1):
        summary_table.rows[idx].cells[0].text = char
        summary_table.rows[idx].cells[1].text = str(count)
        if idx <= highlight_top:
            for paragraph in (
                summary_table.rows[idx].cells[0].paragraphs
                + summary_table.rows[idx].cells[1].paragraphs
            ):
                for run in paragraph.runs:
                    run.bold = True

    output_doc.add_paragraph("\nИсходные таблицы:\n")

    for table_idx, table in enumerate(input_doc.tables):
        output_doc.add_paragraph(f"Таблица {table_idx + 1}:")
        new_table = output_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = "Table Grid"

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    output_doc.save(output_path)


def ask_output_path(default_name: str) -> str | None:
    """Открывает диалог сохранения и возвращает путь либо None."""

    path = filedialog.asksaveasfilename(
        title="Выберите, куда сохранить результат",
        defaultextension=".docx",
        initialfile=default_name,
        filetypes=[("Word файлы", "*.docx")],
    )
    if not path:
        return None

    if Path(path).exists():
        overwrite = messagebox.askyesno(
            "Подтверждение",
            f"Файл {path} уже существует. Перезаписать?",
        )
        if not overwrite:
            return None
    return path


def select_single_file(column_index: int, skip_headers: bool):
    input_file = filedialog.askopenfilename(
        title="Выберите .docx файл",
        filetypes=[("Word файлы", "*.docx")],
    )
    if not input_file:
        return

    output_file = ask_output_path(f"{Path(input_file).stem}_result.docx")
    if not output_file:
        return

    try:
        process_document(
            input_file,
            output_file,
            column_index,
            skip_table_headers=skip_headers,
        )
        messagebox.showinfo(
            "Готово",
            f"Файл обработан и сохранён как:\n{output_file}",
        )
    except Exception as exc:
        messagebox.showerror("Ошибка", f"Произошла ошибка:\n{exc}")


def select_multiple_files(
    column_index: int, status_var: tk.StringVar, skip_headers: bool
):
    files = filedialog.askopenfilenames(
        title="Выберите .docx файлы",
        filetypes=[("Word файлы", "*.docx")],
    )
    if not files:
        return

    target_folder = filedialog.askdirectory(title="Выберите папку для сохранения")
    if not target_folder:
        return

    successes = []
    errors = []
    for idx, file_path in enumerate(files, start=1):
        status_var.set(f"Обработка файла {idx} из {len(files)}...")
        root.update_idletasks()
        destination = Path(target_folder) / f"{Path(file_path).stem}_result.docx"
        if destination.exists():
            overwrite = messagebox.askyesno(
                "Подтверждение",
                f"Файл {destination} уже существует. Перезаписать?",
            )
            if not overwrite:
                errors.append(f"Пропущен {destination} (отменено пользователем)")
                continue

        try:
            process_document(
                file_path,
                destination,
                column_index,
                skip_table_headers=skip_headers,
            )
            successes.append(str(destination))
        except Exception as exc:  # pragma: no cover - GUI feedback
            errors.append(f"{Path(file_path).name}: {exc}")

    status_var.set("")
    if successes:
        messagebox.showinfo(
            "Готово",
            "\n".join(["Сохранены файлы:"] + successes),
        )
    if errors:
        messagebox.showwarning(
            "Ошибки",
            "\n".join(["Не удалось обработать:"] + errors),
        )


root = tk.Tk()
root.title("Обработка Word-файлов")
root.geometry("420x240")

column_var = tk.IntVar(value=2)
skip_headers_var = tk.BooleanVar(value=True)
status_var = tk.StringVar(value="")

description = tk.Label(
    root,
    text=(
        "Выберите столбец с персонажами и обработайте один или несколько .docx файлов."
    ),
    wraplength=380,
    pady=10,
)
description.pack()

column_frame = tk.Frame(root)
column_frame.pack(pady=5)

tk.Label(column_frame, text="Номер столбца (1 = первый):").pack(side=tk.LEFT)
column_spinbox = tk.Spinbox(column_frame, from_=1, to=20, width=5, textvariable=column_var)
column_spinbox.pack(side=tk.LEFT, padx=10)

skip_headers_checkbox = tk.Checkbutton(
    root,
    text="Пропускать первую строку таблиц (заголовки)",
    variable=skip_headers_var,
)
skip_headers_checkbox.pack()

button_frame = tk.Frame(root)
button_frame.pack(pady=15)

single_button = tk.Button(
    button_frame,
    text="Выбрать один файл",
    command=lambda: select_single_file(column_var.get() - 1, skip_headers_var.get()),
    padx=10,
    pady=5,
)
single_button.pack(side=tk.LEFT, padx=5)

multi_button = tk.Button(
    button_frame,
    text="Обработать несколько файлов",
    command=lambda: select_multiple_files(
        column_var.get() - 1, status_var, skip_headers_var.get()
    ),
    padx=10,
    pady=5,
)
multi_button.pack(side=tk.LEFT, padx=5)

status_label = tk.Label(root, textvariable=status_var, fg="gray")
status_label.pack(pady=10)

root.mainloop()
